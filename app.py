import streamlit as st
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="Laudo Ecocardiograma", layout="centered")

# --- Login simples ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    st.title("Login")
    username = st.text_input("Usuário")
    password = st.text_input("Senha", type="password")
    if st.button("Entrar"):
        if username == "admin" and password == "admin":
            st.session_state.logged_in = True
            st.experimental_rerun()
        else:
            st.error("Usuário ou senha incorretos.")
    st.stop()

st.title("Laudo Ecocardiograma")

# --- Dados do paciente ---
with st.form("form_paciente"):
    nome = st.text_input("Nome do paciente", max_chars=50)
    peso = st.number_input("Peso (kg)", min_value=1.0, max_value=300.0, format="%.2f")
    altura = st.number_input("Altura (m)", min_value=30.0, max_value=250.0, format="%.2f")
    genero = st.selectbox("Gênero", ["Masculino", "Feminino"])
    
    st.markdown("### Medidas (em milímetros)")
    aorta = st.number_input("Aorta", min_value=1.0, max_value=100.0, format="%.2f")
    atrio_esq = st.number_input("Átrio Esquerdo", min_value=1.0, max_value=100.0, format="%.2f")
    septo = st.number_input("Septo Interventricular", min_value=1.0, max_value=50.0, format="%.2f")
    parede_post = st.number_input("Parede Posterior", min_value=1.0, max_value=50.0, format="%.2f")
    dd_ve = st.number_input("Diâmetro Diastólico do VE", min_value=10.0, max_value=100.0, format="%.2f")
    ds_ve = st.number_input("Diâmetro Sistólico do VE", min_value=5.0, max_value=80.0, format="%.2f")
    vol_atrio_esq = st.number_input("Volume do Átrio Esquerdo (mL)", min_value=1.0, max_value=300.0, format="%.2f")

    submitted = st.form_submit_button("Calcular")

if submitted:
    # Ajuste: altura em cm convertida para metros para fórmula de Dubois
    altura_m = altura * 100  # cm -> m

    # Cálculo da área de superfície corporal (Dubois)
    bsa = 0.007184 * (altura ** 0.725) * (peso ** 0.425)  # altura em cm, peso em kg

    # Cálculo do índice de massa do VE (IMVE) em g/m²
    # Fórmula do volume do VE (Devereux): 0.8 * [1.04 * ((DD + Sep + PP)^3 - (DD)^3)] + 0.6 g
    massa_ve = 0.8 * (1.04 * (((dd_ve + septo + parede_post) ** 3) - (dd_ve ** 3))) + 0.6
    imve = massa_ve / bsa

    # Espessura relativa (ER) = (2 x PP) / DD
    er = (2 * parede_post) / dd_ve

    # Volume do átrio esquerdo indexado
    vol_atrio_esq_i = vol_atrio_esq / bsa

    # Fração de ejeção pelo método de Teicholz
    fe = ((dd_ve ** 3) - (ds_ve ** 3)) / (dd_ve ** 3) * 100

    # Exibição resultados
    st.markdown(f"**Área de Superfície Corporal (ASC):** {bsa:.2f} m²")
    st.markdown(f"**Índice de Massa do VE:** {imve:.2f} g/m²")
    st.markdown(f"**Espessura Relativa:** {er:.2f}")
    st.markdown(f"**Volume do Átrio Esquerdo Indexado:** {vol_atrio_esq_i:.2f} mL/m²")
    st.markdown(f"**Fração de Ejeção (Teicholz):** {fe:.2f} %")

    # --- Geração automática do texto do laudo ---
    laudo = []

    # Ventrículo esquerdo
    if imve < 115:
        laudo.append("Ventrículo esquerdo: cavidade com dimensões normais, paredes com espessura normal, ausência de alteração da contratilidade segmentar, função sistólica e diastólica normal.")
    else:
        laudo.append("Ventrículo esquerdo: hipertrofia ventricular esquerda evidenciada, com espessura aumentada das paredes.")

    # Átrio esquerdo
    if vol_atrio_esq_i < 34:
        laudo.append("Átrio esquerdo: dimensões normais, volume indexado normal.")
    else:
        laudo.append("Átrio esquerdo: aumento do volume indexado do átrio esquerdo.")

    # VD e AD - normal padrão pois não temos medidas específicas
    laudo.append("Ventrículo direito e átrio direito: cavidades com dimensões normais, função sistólica normal.")

    # Valvas - normal padrão
    laudo.append("Valva aórtica: folhetos com espessura e mobilidade normais, abertura valvar preservada, ausência de refluxo aórtico.")
    laudo.append("Valva mitral: folhetos com espessura e mobilidade normais, abertura valvar preservada, refluxo fisiológico.")
    laudo.append("Valva tricúspide: folhetos com espessura e mobilidade normais, refluxo mínimo, fisiológico.")
    laudo.append("Valva pulmonar: folhetos com espessura e mobilidade normais, refluxo mínimo, fisiológico.")

    # Pericárdio
    laudo.append("Pericárdio: ausência de derrame pericárdico.")

    # Conclusão
    if imve < 115 and vol_atrio_esq_i < 34 and fe > 55:
        laudo.append("Conclusão: Ecocardiograma transtorácico normal.")
    else:
        laudo.append("Conclusão: Alterações ecocardiográficas detectadas, correlacionar clinicamente.")

    texto_laudo = "\n\n".join(laudo)

    st.subheader("Texto do Laudo")
    texto_editavel = st.text_area("Você pode editar o texto abaixo antes de gerar o arquivo", texto_laudo, height=350)

    # --- Criação do arquivo DOCX para download ---
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    document.add_heading('Laudo Ecocardiograma', level=1)
    document.add_paragraph(f"Paciente: {nome}")
    document.add_paragraph(f"Peso: {peso:.2f} kg")
    document.add_paragraph(f"Altura: {altura:.2f} cm")
    document.add_paragraph(f"Gênero: {genero}")
    document.add_paragraph(f"Área de Superfície Corporal (ASC): {bsa:.2f} m²\n")

    for paragrafo in texto_editavel.strip().split('\n\n'):
        document.add_paragraph(paragrafo.strip())

    # Salvar em buffer para download (evita /tmp que pode dar erro em Windows)
    import io
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📄 Baixar Laudo em DOCX",
        data=buffer,
        file_name="laudo_ecocardiograma.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
